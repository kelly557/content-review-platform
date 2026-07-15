"""Document text extraction for audit point parsing."""
from __future__ import annotations

import io
from pathlib import Path

from pypdf import PdfReader


def extract_text_from_pdf(content: bytes) -> str:
    """Extract text from PDF file."""
    reader = PdfReader(io.BytesIO(content))
    texts = []
    for page in reader.pages:
        text = page.extract_text()
        if text:
            texts.append(text)
    return "\n".join(texts)


def extract_text_from_docx(content: bytes) -> str:
    """Extract text from DOCX file."""
    from docx import Document

    doc = Document(io.BytesIO(content))
    texts = []
    for para in doc.paragraphs:
        if para.text.strip():
            texts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                texts.append(row_text)
    return "\n".join(texts)


def extract_text_from_file(content: bytes, filename: str) -> str:
    """Extract text from supported file types."""
    ext = Path(filename).suffix.lower()

    if ext == ".pdf":
        return extract_text_from_pdf(content)
    elif ext in (".doc", ".docx"):
        return extract_text_from_docx(content)
    elif ext in (".txt", ".md"):
        return content.decode("utf-8", errors="replace")
    else:
        raise ValueError(f"不支持的文件格式: {ext}")
